import matplotlib.pyplot as plt; plt.rcdefaults()
import numpy as np
import matplotlib.pyplot as plt
import os
import json
import sys
from colorama import init, Fore
from PIL import Image
import matplotlib.image as mpimg
import random	
import math
from docx import Document
import boto3
import botocore


init() #colorama


#Test results will contain tuples of the tests results rom the ingest engine
#_TestResultGroup = []

#Contain all of the organized groups
_TestGroups = []
_decodedJson = dict()
_localDir = ""

_jsonPath = ""
_writePath = ""

_imagePaths = []

BUCKETNAME = "doctordocx-patientqueue"


_count = 0.0
_totalDotsH = 0.0

DEBUG = False
SHOWCHART = False
BOGUSVALS = True
FIRSTPAGE = True
DONE = False
AUTOPATH = False
PAGECOUNT = 1
GENDOCX = True

def main():
	global DONE 
	global _jsonPath
	global _writePath

	# if len(sys.argv) < 3:
	# 	_jsonPath = sys.argv[1]
	# 	_writePath = sys.argv[2]

	print(Fore.GREEN + "Generating Table")
	ingestDataFile()
	#generateChart()
	#saveChart()
	print(Fore.BLUE + "Generating Image Page	")
	#loadNormalCurve()
	#concatImages()
	#generateDoc()

def list_bucket_objects(bucket_name):
    """List the objects in an Amazon S3 bucket

    :param bucket_name: string
    :return: List of bucket objects. If error, return None.
    """

    # Retrieve the list of bucket objects
    s_3 = boto3.client('s3')
    try:
        response = s_3.list_objects_v2(Bucket=BUCKETNAME)
    except ClientError as e:
        # AllAccessDisabled error == bucket not found
        logging.error(e)
        return None

    # Only return the contents if we found some keys
    if response['KeyCount'] > 0:
        return response['Contents']

    return None


#Assume the data file is in the same directory with name dataFile.txt
def ingestDataFile():
	print("Intaking Data File")
	global DEBUG
	global _decodedJson
	global _localDir
	global AUTOPATH

	_localDir = os.getcwd()

	s3 = boto3.client('s3')

	if DEBUG:
		print("Local Directory is: " + _localDir)

    objects = list_bucket_objects(bucket_name)


	if objects is not None:
        # List the object names
		s3.download_file(BUCKETNAME, objects[0], "local.json")

    else:
        return 0

	_dataFile = open('local.json', "r")

	_rawDataJson = _dataFile.read()
	_decodedJson = json.loads(_rawDataJson)
	print(_decodedJson.pop("ATTENTION").pop("W4-DSFd"))

	_dataFile.close()

def generateChart():
	global DEBUG
	global _decodedJson
	global SHOWCHART
	global _count
	global BOGUSVALS
	global PAGECOUNT

	print("Generating Chart")


	_categories = []
	for _key in _decodedJson.keys():
		_categories.append(_key)

	if DEBUG:
		print("Categories: ")
		print(_categories)

	_xList = []
	_yList = []
	_colors = []

	for _cat in _categories:
		_xList.append(r"$\bf{" + str(_cat.replace(" ", "\ ")) + "}$") #magic, don't touch
		_yList.append(0)
		_colors.append('black')

		# _testList = []
		# for _key in _decodedJson[_cat].keys():
		# 	_testList.append(_key)

		for index in range(0, len(_decodedJson[_cat])):
			_testName = _decodedJson[_cat][index]['RelatedTest']['Name']
			_testVal = _decodedJson[_cat][index]['ZScore']
			if BOGUSVALS:
				_testVal = random.uniform(-5,5)
			_xList.append(_testName)
			_yList.append(_testVal)
			_colors.append(getColor(_testVal))


		_xList.append("")
		_yList.append(0)
		_colors.append('black')
	
	_xList.pop()
	_yList.pop()
	_colors.pop()

	while not len(_xList) % 36 == 0:
		_xList.append("")
		_yList.append(0)
		_colors.append('black')

	if DEBUG:
		print(_xList)
		print(_yList)

	#_xList.reverse()
	#_yList.reverse()
	#_colors.reverse()

	PAGECOUNT = math.ceil(len(_xList) / 36.0)
	for page in range(1, PAGECOUNT + 1):
		_tmpX = _xList[0 + 36 * (page - 1): 36 * page]
		_tmpY = _yList[0 + 36 * (page - 1): 36 * page]
		_tmpcolors = _colors[0 + 36 * (page - 1): 36 * page]

		_tmpX.reverse()
		_tmpY.reverse()
		_tmpcolors.reverse()

		_count = len(_tmpX)
		_yScal = np.arange(_count)
		plt.barh(_yScal, _tmpY, color = _tmpcolors, align='center', alpha=1)
		plt.yticks(_yScal, _tmpX, rotation = 35, fontsize = 10)
		plt.xlim([-4, 4])
		plt.grid(b=True, axis = 'x')
		plt.subplots_adjust(left = 0.25)
		#plt.tight_layout(pad = 0.5)
	#	if SHOWCHART:
	#		plt.show()
		saveChart(page)

def getColor(_val):
	#color = 'green'

	if _val < 0 and _val >= -1:
		color = 'yellow'
	elif _val < -1 and _val >= -2:
		color = 'orange'
	elif _val < -2: 
		color = 'red'
	else: 
		color = 'green'

	return color


def saveChart(pageNumber):
	global _localDir

	print(Fore.YELLOW + "Saving Graph " + str(pageNumber))
	fig = plt.gcf()
	fig.set_size_inches(8, 10.5)
	fig.savefig(_localDir + "\\graph" + str(pageNumber) + ".png", tranparent=True, dpi = 300, orientation = 'portrait', pad_inches = 0)
	plt.clf()

def loadNormalCurve():
	print("Loading Curve Image")
	img = mpimg.imread(_localDir + "\\NormalCurve.png")

def concatImages():
	global PAGECOUNT
	global _imagePaths

	for _page in range(1, PAGECOUNT + 1):
		print(Fore.CYAN + "Generating Visualization Page " + str(_page))
		_curve = Image.open(_localDir + "\\NormalCurveResized.png")
		_table = Image.open(_localDir + "\\graph" + str(_page) + ".png")

		_dest = Image.new('RGB', (_table.width, _table.height + _curve.height), (255, 255, 255))
		_dest.paste(_table, (0, _curve.height - 470))
		_dest.paste(_curve, (186, 0))

		_dest = _dest.crop((0, 0, _dest.width, 300 * 10))
		_dest.save(_localDir + "\\renderedVisualization" + str(_page) + ".png")
		_imagePaths.append(_localDir + "\\renderedVisualization" + str(_page) + ".png")
	# print("PRINTING PAGE: " + str(PAGECOUNT))

	# _curve = Image.open(_localDir + "\\NormalCurveResized.png")
	# _table = Image.open(_localDir + "\\graph" + str(PAGECOUNT) + ".png")

	# _dest = Image.new('RGB', (_table.width, _table.height + _curve.height), (255, 255, 255))

	# _dest.paste(_table, (0, _curve.height - 150))
	# _dest.paste(_curve, (55, 0))

	# _dest = _dest.crop((0, 0, _dest.width, 300 * 10))
	# _dest.save(_localDir + "\\finishedVisualization" + str(PAGECOUNT) + ".png")

	# _cutoff = _table.height - (3000 - 601 - 150)
	# print(_cutoff)
	# if _cutoff <= 0: 
	# 	DONE = True
	# else:
	# 	_table = _table.crop((0, _cutoff, _table.width, _table.height - _cutoff))
	# 	_table.save(_localDir + "\\graph" + str(PAGECOUNT + 1) + ".png")

	# FIRSTPAGE = False
	# PAGECOUNT = PAGECOUNT + 1

def generateDoc():
	document = Document()

	para = document.add_paragraph()
	run = para.add_run()

	for image in _imagePaths:
		run.add_picture(image)

	document.save(_localDir + "/visualization.docx")

if __name__ == "__main__":
    main()